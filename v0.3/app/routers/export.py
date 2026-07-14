"""Exporte.

PDF (WeasyPrint):
  /export/wettkampf/{wid}/startliste.pdf
  /export/wettkampf/{wid}/ergebnisse.pdf
  /export/wettkampf/{wid}/urkunden.pdf
  /export/wettkampf/{wid}/wertungskarten.pdf
  /export/tag/{tid}/ergebnisse.pdf

Siegerliste / Office:
  /export/wettkampf/{wid}/siegerliste.xlsx   (Excel)
  /export/wettkampf/{wid}/siegerliste.csv    (CSV)
  /export/tag/{tid}/siegerliste.xlsx         (Excel, ein Blatt pro Wettkampf)
  /export/wettkampf/{wid}/urkunden.docx      (Word, eine Seite pro Athlet,
                                              vor dem Druck frei anpassbar)
"""
import base64
import csv
import io
import json
from typing import Optional

from fastapi import APIRouter, Depends, Query, Request
from fastapi.responses import Response, RedirectResponse, StreamingResponse
from sqlalchemy.orm import Session
from weasyprint import HTML

from auth import require_user
from database import get_db
from models import (
    Wettkampf, WettkampfTag, PersonenHasWettkampf, GeraeteHasWettkampf,
)
from scoring import get_strategy
from services.rangliste import (
    einzel_rangliste, einzel_rangliste_mit_geraeten, mannschaft_rangliste,
)
from services.backup import snapshot_tag
from views import templates as jinja

router = APIRouter(prefix="/export")


def _logo_data_url(blob: bytes | None, mime: str | None) -> str | None:
    if not blob:
        return None
    return f"data:{mime or 'image/png'};base64,{base64.b64encode(blob).decode()}"


def _render_pdf(template_name: str, **ctx) -> bytes:
    tpl = jinja.env.get_template(template_name)
    html = tpl.render(**ctx)
    return HTML(string=html, base_url=".").write_pdf()


def _pdf_response(data: bytes, filename: str) -> Response:
    return Response(
        content=data,
        media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename="{filename}"'},
    )


@router.get("/wettkampf/{wid}/startliste.pdf")
def startliste(wid: int, db: Session = Depends(get_db),
               user=Depends(require_user())):
    wk = db.get(Wettkampf, wid)
    if not wk:
        return RedirectResponse("/tage", status_code=303)
    starter = (
        db.query(PersonenHasWettkampf)
        .filter_by(Wettkampf_id=wid)
        .order_by(
            PersonenHasWettkampf.Riege_id.is_(None),
            PersonenHasWettkampf.Riege_id,
            PersonenHasWettkampf.Startnummer,
        )
        .all()
    )
    pdf = _render_pdf("pdf/startliste.html", wk=wk, starter=starter)
    return _pdf_response(pdf, f"startliste-wk{wid}.pdf")


@router.get("/wettkampf/{wid}/ergebnisse.pdf")
def ergebnisse(wid: int,
               top: Optional[int] = Query(None, ge=1),
               db: Session = Depends(get_db),
               user=Depends(require_user())):
    wk = db.get(Wettkampf, wid)
    if not wk:
        return RedirectResponse("/tage", status_code=303)
    einzel, geraete = einzel_rangliste_mit_geraeten(db, wid)
    if top:
        einzel = [r for r in einzel if r["Platz"] <= top]
    teams = mannschaft_rangliste(db, wid, wk.Mannschaft_Groesse) if wk.Typ != "Einzel" else []
    if top:
        teams = [t for t in teams if t["Platz"] <= top]
    pdf = _render_pdf("pdf/ergebnisse.html", wk=wk, einzel=einzel,
                      teams=teams, geraete=geraete, top=top)
    suffix = f"-top{top}" if top else ""
    return _pdf_response(pdf, f"ergebnisse-wk{wid}{suffix}.pdf")


@router.get("/wettkampf/{wid}/urkunden.pdf")
def urkunden(wid: int,
             top: Optional[int] = Query(None, ge=1),
             db: Session = Depends(get_db),
             user=Depends(require_user())):
    wk = db.get(Wettkampf, wid)
    if not wk:
        return RedirectResponse("/tage", status_code=303)
    einzel = einzel_rangliste(db, wid)
    if top:
        einzel = [r for r in einzel if r["Platz"] <= top]
    # nur Athleten mit echtem Score
    einzel = [r for r in einzel if (r.get("GesamtScore") or 0) > 0]
    logo_url = _logo_data_url(wk.tag.Logo, wk.tag.Logo_MimeType)
    pdf = _render_pdf("pdf/urkunden.html", wk=wk, einzel=einzel, logo_url=logo_url)
    suffix = f"-top{top}" if top else ""
    return _pdf_response(pdf, f"urkunden-wk{wid}{suffix}.pdf")


@router.get("/tag/{tid}/ergebnisse.pdf")
def tag_ergebnisse(tid: int, db: Session = Depends(get_db),
                   user=Depends(require_user())):
    tag = db.get(WettkampfTag, tid)
    if not tag:
        return RedirectResponse("/tage", status_code=303)
    blocks = []
    for w in tag.wettkaempfe:
        blocks.append({
            "wk": w,
            "einzel": einzel_rangliste(db, w.idWettkampf),
            "teams": mannschaft_rangliste(db, w.idWettkampf, w.Mannschaft_Groesse) if w.Typ != "Einzel" else [],
        })
    pdf = _render_pdf("pdf/tag_ergebnisse.html", tag=tag, blocks=blocks)
    return _pdf_response(pdf, f"ergebnisse-tag{tid}.pdf")


# F3: Wertungskarten als PDF (leer, fuer Kampfrichter zum Ausfuellen).
# Mit ?geraet_id: eine Karte pro Athlet+Versuch fuer DIESES Geraet.
# Ohne ?geraet_id: Laufkarten — eine Seite pro Person mit allen Geraeten drauf.
@router.get("/wettkampf/{wid}/wertungskarten.pdf")
def wertungskarten(wid: int, geraet_id: Optional[int] = Query(None),
                   db: Session = Depends(get_db),
                   user=Depends(require_user())):
    wk = db.get(Wettkampf, wid)
    if not wk:
        return RedirectResponse("/tage", status_code=303)
    geraete_zu = sorted(wk.geraete_zuordnung, key=lambda g: g.Reihenfolge)
    if geraet_id:
        geraete_zu = [g for g in geraete_zu if g.Geraete_id == geraet_id]
    if not geraete_zu:
        return RedirectResponse(f"/wettkampf/{wid}", status_code=303)

    starter = (
        db.query(PersonenHasWettkampf)
        .filter_by(Wettkampf_id=wid)
        .order_by(
            PersonenHasWettkampf.Riege_id.is_(None),
            PersonenHasWettkampf.Riege_id,
            PersonenHasWettkampf.Startnummer,
        ).all()
    )

    geraete_blocks = [
        {"ghw": g, "kriterien": get_strategy(g.berechnung.Regel_Kuerzel).alle_kriterien}
        for g in geraete_zu
    ]
    if geraet_id:
        pdf = _render_pdf("pdf/wertungskarten.html",
                          wk=wk, starter=starter, geraete_blocks=geraete_blocks)
        fn = f"wertungskarten-wk{wid}-g{geraet_id}.pdf"
    else:
        pdf = _render_pdf("pdf/laufkarten.html",
                          wk=wk, starter=starter, geraete_blocks=geraete_blocks)
        fn = f"laufkarten-wk{wid}.pdf"
    return _pdf_response(pdf, fn)


# F4: CSV-Export der Ergebnisse
@router.get("/wettkampf/{wid}/ergebnisse.csv")
def ergebnisse_csv(wid: int, db: Session = Depends(get_db),
                   user=Depends(require_user())):
    wk = db.get(Wettkampf, wid)
    if not wk:
        return RedirectResponse("/tage", status_code=303)
    einzel, geraete = einzel_rangliste_mit_geraeten(db, wid)
    buf = io.StringIO()
    w = csv.writer(buf, delimiter=";", lineterminator="\n")
    w.writerow(["Platz", "Vorname", "Nachname", "Verein"]
               + [g.Name for g in geraete] + ["Gesamt"])
    for r in einzel:
        row = [r["Platz"], r["Vorname"], r["Nachname"], r.get("Verein_Kuerzel") or ""]
        for g in geraete:
            s = r["geraete_scores"].get(g.idGeraete)
            row.append(f"{s:.3f}" if s is not None else "")
        row.append(f"{float(r['GesamtScore']):.3f}")
        w.writerow(row)
    return StreamingResponse(
        iter([buf.getvalue().encode("utf-8-sig")]),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="ergebnisse-wk{wid}.csv"'},
    )


# ---------------------------------------------------------------------------
# Siegerlisten (Excel / CSV) + Urkunden (Word)
# ---------------------------------------------------------------------------

def _siegerliste_rows(db: Session, wid: int, top: Optional[int]):
    """Gemeinsame Datenbasis fuer Siegerlisten-Exporte."""
    einzel, geraete = einzel_rangliste_mit_geraeten(db, wid)
    if top:
        einzel = [r for r in einzel if r["Platz"] <= top]
    return einzel, geraete


def _xlsx_sheet_fuellen(ws, wk, einzel, geraete, teams):
    """Schreibt eine Siegerliste in ein openpyxl-Worksheet."""
    from openpyxl.styles import Alignment, Font

    bold = Font(bold=True)
    title = Font(bold=True, size=14)

    ws.append([f"{wk.Wettkampf_Nr} · {wk.Name}"])
    ws.cell(row=1, column=1).font = title
    ws.append([f"{wk.tag.Name} · {wk.tag.Wettkampf_Datum} · "
               f"{wk.altersklasse.Kuerzel} ({wk.altersklasse.Bezeichnung})"])
    ws.append([])

    header = ["Platz", "Nachname", "Vorname", "Verein"] + [g.Name for g in geraete] + ["Gesamt"]
    ws.append(header)
    header_row = ws.max_row
    for col in range(1, len(header) + 1):
        ws.cell(row=header_row, column=col).font = bold

    for r in einzel:
        row = [r["Platz"], r["Nachname"], r["Vorname"], r.get("Verein_Kuerzel") or ""]
        for g in geraete:
            s = r["geraete_scores"].get(g.idGeraete)
            row.append(round(float(s), 3) if s is not None else None)
        row.append(round(float(r["GesamtScore"]), 3))
        ws.append(row)

    if teams:
        ws.append([])
        ws.append(["Mannschaften"])
        ws.cell(row=ws.max_row, column=1).font = bold
        ws.append(["Platz", "Mannschaft", "Mitglieder", "Gesamt"])
        hr = ws.max_row
        for col in range(1, 5):
            ws.cell(row=hr, column=col).font = bold
        for t in teams:
            ws.append([t["Platz"], t["Mannschaft_Name"], t["Mitglieder_Gesamt"],
                       round(float(t["GesamtScore"]), 3)])

    # Spaltenbreiten grob anpassen
    widths = [7, 20, 16, 10] + [12] * len(geraete) + [10]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[ws.cell(row=header_row, column=i).column_letter].width = w
    ws.cell(row=header_row, column=1).alignment = Alignment(horizontal="left")


def _xlsx_response(wb, filename: str) -> Response:
    buf = io.BytesIO()
    wb.save(buf)
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/wettkampf/{wid}/siegerliste.xlsx")
def siegerliste_xlsx(wid: int,
                     top: Optional[int] = Query(None, ge=1),
                     db: Session = Depends(get_db),
                     user=Depends(require_user())):
    from openpyxl import Workbook
    wk = db.get(Wettkampf, wid)
    if not wk:
        return RedirectResponse("/tage", status_code=303)
    einzel, geraete = _siegerliste_rows(db, wid, top)
    teams = mannschaft_rangliste(db, wid, wk.Mannschaft_Groesse) if wk.Typ != "Einzel" else []
    if top:
        teams = [t for t in teams if t["Platz"] <= top]
    wb = Workbook()
    _xlsx_sheet_fuellen(wb.active, wk, einzel, geraete, teams)
    wb.active.title = f"WK {wk.Wettkampf_Nr}"[:31]
    suffix = f"-top{top}" if top else ""
    return _xlsx_response(wb, f"siegerliste-wk{wid}{suffix}.xlsx")


@router.get("/tag/{tid}/siegerliste.xlsx")
def tag_siegerliste_xlsx(tid: int, db: Session = Depends(get_db),
                         user=Depends(require_user())):
    """Ein Excel fuer den ganzen Wettkampftag — ein Blatt pro Wettkampf."""
    from openpyxl import Workbook
    tag = db.get(WettkampfTag, tid)
    if not tag:
        return RedirectResponse("/tage", status_code=303)
    wb = Workbook()
    wb.remove(wb.active)
    for w in sorted(tag.wettkaempfe, key=lambda x: x.Wettkampf_Nr):
        einzel, geraete = _siegerliste_rows(db, w.idWettkampf, None)
        teams = (mannschaft_rangliste(db, w.idWettkampf, w.Mannschaft_Groesse)
                 if w.Typ != "Einzel" else [])
        ws = wb.create_sheet(title=f"WK {w.Wettkampf_Nr}"[:31])
        _xlsx_sheet_fuellen(ws, w, einzel, geraete, teams)
    if not wb.sheetnames:
        wb.create_sheet(title="leer")
    return _xlsx_response(wb, f"siegerliste-tag{tid}.xlsx")


@router.get("/wettkampf/{wid}/siegerliste.csv")
def siegerliste_csv(wid: int,
                    top: Optional[int] = Query(None, ge=1),
                    db: Session = Depends(get_db),
                    user=Depends(require_user())):
    wk = db.get(Wettkampf, wid)
    if not wk:
        return RedirectResponse("/tage", status_code=303)
    einzel, geraete = _siegerliste_rows(db, wid, top)
    buf = io.StringIO()
    w = csv.writer(buf, delimiter=";", lineterminator="\n")
    w.writerow(["Platz", "Nachname", "Vorname", "Verein"]
               + [g.Name for g in geraete] + ["Gesamt"])
    for r in einzel:
        row = [r["Platz"], r["Nachname"], r["Vorname"], r.get("Verein_Kuerzel") or ""]
        for g in geraete:
            s = r["geraete_scores"].get(g.idGeraete)
            row.append(f"{s:.3f}" if s is not None else "")
        row.append(f"{float(r['GesamtScore']):.3f}")
        w.writerow(row)
    suffix = f"-top{top}" if top else ""
    return StreamingResponse(
        iter([buf.getvalue().encode("utf-8-sig")]),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition":
                 f'attachment; filename="siegerliste-wk{wid}{suffix}.csv"'},
    )


@router.get("/wettkampf/{wid}/urkunden.docx")
def urkunden_docx(wid: int,
                  top: Optional[int] = Query(None, ge=1),
                  db: Session = Depends(get_db),
                  user=Depends(require_user())):
    """Urkunden als Word-Dokument: eine Seite pro platziertem Athlet.
    Kann vor dem Druck in Word beliebig angepasst werden (Schrift, Logo, Text)."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.shared import Cm, Pt

    wk = db.get(Wettkampf, wid)
    if not wk:
        return RedirectResponse("/tage", status_code=303)
    einzel = einzel_rangliste(db, wid)
    if top:
        einzel = [r for r in einzel if r["Platz"] <= top]
    einzel = [r for r in einzel if (r.get("GesamtScore") or 0) > 0]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)

    def zeile(text: str, size: int, bold: bool = False, space_after: int = 12):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        return p

    datum = wk.tag.Wettkampf_Datum.strftime("%d.%m.%Y")
    for i, r in enumerate(einzel):
        if i > 0:
            doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
        if wk.tag.Logo:
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(io.BytesIO(wk.tag.Logo), height=Cm(3))
            except Exception:
                pass  # unlesbares Logo-Format soll den Export nicht stoppen
        zeile(wk.tag.Name, 16, space_after=24)
        zeile("URKUNDE", 40, bold=True, space_after=30)
        zeile(f"{wk.Name} · {wk.altersklasse.Bezeichnung}", 14, space_after=30)
        zeile(f"{r['Vorname']} {r['Nachname']}", 28, bold=True, space_after=6)
        verein = r.get("Verein_Name") or r.get("Verein_Kuerzel") or ""
        if verein:
            zeile(verein, 14, space_after=24)
        zeile(f"{r['Platz']}. Platz", 24, bold=True, space_after=12)
        zeile(f"mit {float(r['GesamtScore']):.3f} Punkten", 14, space_after=48)
        ort = f"{wk.tag.Ort}, " if wk.tag.Ort else ""
        zeile(f"{ort}{datum}", 12, space_after=0)

    buf = io.BytesIO()
    doc.save(buf)
    suffix = f"-top{top}" if top else ""
    return Response(
        content=buf.getvalue(),
        media_type=("application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document"),
        headers={"Content-Disposition":
                 f'attachment; filename="urkunden-wk{wid}{suffix}.docx"'},
    )


# F2: JSON-Snapshot eines Wettkampftags (Backup)
@router.get("/tag/{tid}/backup.json")
def backup_json(tid: int, db: Session = Depends(get_db),
                user=Depends(require_user("admin"))):
    data = snapshot_tag(db, tid)
    if not data:
        return RedirectResponse("/tage", status_code=303)
    payload = json.dumps(data, indent=2, ensure_ascii=False, default=str)
    return Response(
        content=payload, media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="backup-tag{tid}.json"'},
    )
